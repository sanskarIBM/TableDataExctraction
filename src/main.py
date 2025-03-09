import os
import json
import logging
import argparse
import time
from typing import List, Dict, Any, Tuple, Optional

# For PDF processing
import fitz  # PyMuPDF
import cv2
import numpy as np
import pytesseract
from PIL import Image

# For DOCX processing
import docx
from docx.table import Table as DocxTable

# For table detection and extraction
import pandas as pd
from dotenv import load_dotenv
import groq

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[logging.StreamHandler()]
)
logger = logging.getLogger(__name__)

load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
client = groq.Client(api_key=GROQ_API_KEY) if GROQ_API_KEY else None


class DocumentAgent:
    """Agent responsible for identifying tables in documents."""

    def __init__(self, file_path: str):
        self.file_path = file_path
        self.file_extension = os.path.splitext(file_path)[1].lower()
        self.tables_info = []

    def process(self) -> List[Dict[str, Any]]:
        """Process the document and identify tables."""
        if self.file_extension == '.pdf':
            return self._process_pdf()
        elif self.file_extension in ['.docx', '.doc']:
            return self._process_docx()
        elif self.file_extension in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
            return self._process_image()
        else:
            raise ValueError(f"Unsupported file format: {self.file_extension}")

    def _process_pdf(self) -> List[Dict[str, Any]]:
        """Identify tables in PDF documents using multiple methods."""
        logger.info(f"Identifying tables in PDF: {self.file_path}")

        tables_tabula = []
        tables_camelot = []
        tables_visual = []

        # Method 1: Try tabula-py with safer error handling
        try:
            # Dynamically import tabula to avoid initial import errors
            from tabula import read_pdf

            # Use explicit encoding parameter and handle potential errors
            try:
                dfs_default = read_pdf(self.file_path, pages='all', multiple_tables=True, encoding='latin1')
                dfs_lattice = read_pdf(self.file_path, pages='all', multiple_tables=True, lattice=True,
                                       encoding='latin1')
                dfs_stream = read_pdf(self.file_path, pages='all', multiple_tables=True, stream=True, encoding='latin1')

                # Combine all results
                dfs = dfs_default + dfs_lattice + dfs_stream

                for i, df in enumerate(dfs):
                    if not df.empty and len(df.columns) > 1:
                        tables_tabula.append({
                            'method': 'tabula',
                            'page': i % 10 + 1,  # Rough estimation
                            'data': df
                        })
            except Exception as e:
                logger.warning(f"Tabula extraction error with encoding 'latin1': {e}")
                # Try with a different encoding
                try:
                    dfs = read_pdf(self.file_path, pages='all', multiple_tables=True, encoding='utf-8')
                    for i, df in enumerate(dfs):
                        if not df.empty and len(df.columns) > 1:
                            tables_tabula.append({
                                'method': 'tabula',
                                'page': i % 10 + 1,
                                'data': df
                            })
                except Exception as e2:
                    logger.warning(f"Tabula extraction error with encoding 'utf-8': {e2}")
        except ImportError:
            logger.warning("Tabula not available, skipping this method")

        # Method 2: Use camelot with safer error handling
        try:
            import camelot
            try:
                tables = []
                try:
                    tables.extend(camelot.read_pdf(self.file_path, pages='all', flavor='lattice'))
                except Exception as e:
                    logger.warning(f"Camelot lattice flavor error: {e}")

                try:
                    tables.extend(camelot.read_pdf(self.file_path, pages='all', flavor='stream', edge_tol=500))
                except Exception as e:
                    logger.warning(f"Camelot stream flavor error: {e}")

                for i, table in enumerate(tables):
                    if not table.df.empty:
                        tables_camelot.append({
                            'method': 'camelot',
                            'page': table.page,
                            'accuracy': table.accuracy,
                            'data': table.df
                        })
            except Exception as e:
                logger.warning(f"Camelot extraction error: {e}")
        except ImportError:
            logger.warning("Camelot not available, skipping this method")

        # Method 3: Visual detection with improved error handling
        try:
            doc = fitz.open(self.file_path)
            for page_num, page in enumerate(doc):
                try:
                    pix = page.get_pixmap(alpha=False)
                    img_array = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.h, pix.w, pix.n)

                    # Convert to grayscale if needed
                    if pix.n >= 3:
                        gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
                    else:
                        gray = img_array

                    # Apply binary threshold
                    _, binary = cv2.threshold(gray, 200, 255, cv2.THRESH_BINARY_INV)

                    # Detect horizontal and vertical lines
                    horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (50, 1))
                    vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 50))

                    horizontal_lines = cv2.erode(binary, horizontal_kernel, iterations=1)
                    horizontal_lines = cv2.dilate(horizontal_lines, horizontal_kernel, iterations=3)

                    vertical_lines = cv2.erode(binary, vertical_kernel, iterations=1)
                    vertical_lines = cv2.dilate(vertical_lines, vertical_kernel, iterations=3)

                    # Combine lines
                    table_mask = cv2.bitwise_or(horizontal_lines, vertical_lines)

                    # Find contours
                    contours, _ = cv2.findContours(table_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

                    # Filter potential tables
                    for contour in contours:
                        area = cv2.contourArea(contour)
                        if area > 5000:
                            x, y, w, h = cv2.boundingRect(contour)
                            if 0.2 < w / h < 10 and w > 100 and h > 100:
                                tables_visual.append({
                                    'method': 'visual',
                                    'page': page_num + 1,
                                    'bbox': (x, y, x + w, y + h),
                                    'raw_image': img_array[y:y + h, x:x + w].copy()  # Ensure a copy
                                })
                except Exception as e:
                    logger.warning(f"Error processing page {page_num + 1}: {e}")
            doc.close()
        except Exception as e:
            logger.warning(f"Visual detection error: {e}")

        # Fallback: Extract page text directly if no tables detected
        all_tables = tables_tabula + tables_camelot + tables_visual

        if not all_tables:
            logger.info("No tables detected with standard methods, trying text extraction")
            try:
                doc = fitz.open(self.file_path)
                for page_num, page in enumerate(doc):
                    text = page.get_text()
                    if text.strip():
                        # Create structured data from text
                        text_lines = [line.strip() for line in text.split('\n') if line.strip()]
                        if len(text_lines) > 5:  # Minimum number of lines for potential table
                            all_tables.append({
                                'method': 'text',
                                'page': page_num + 1,
                                'text_content': text_lines
                            })
                doc.close()
            except Exception as e:
                logger.warning(f"Text extraction error: {e}")

            # If still no tables, use full-page images as last resort
            if not all_tables:
                logger.info("No text-based tables found, using full-page images")
                try:
                    doc = fitz.open(self.file_path)
                    for page_num, page in enumerate(doc):
                        pix = page.get_pixmap(alpha=False)
                        img_array = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.h, pix.w, pix.n)

                        all_tables.append({
                            'method': 'fullpage',
                            'page': page_num + 1,
                            'raw_image': img_array.copy()  # Ensure a copy
                        })
                    doc.close()
                except Exception as e:
                    logger.warning(f"Full-page extraction error: {e}")

        # Sort by page number
        all_tables.sort(key=lambda x: x.get('page', 0))
        return all_tables

    def _process_docx(self) -> List[Dict[str, Any]]:
        """Identify tables in DOCX documents."""
        logger.info(f"Identifying tables in DOCX: {self.file_path}")

        tables = []
        try:
            doc = docx.Document(self.file_path)
            for i, table in enumerate(doc.tables):
                tables.append({
                    'method': 'docx',
                    'index': i,
                    'rows': len(table.rows),
                    'cols': len(table.columns),
                    'table_obj': table
                })

            # Try mammoth if no tables found
            if len(tables) == 0:
                try:
                    import mammoth
                    result = mammoth.convert_to_html(self.file_path)
                    html = result.value

                    if "<table" in html:
                        logger.info("Found tables using mammoth HTML conversion")
                        tables.append({
                            'method': 'mammoth',
                            'html': html
                        })
                except ImportError:
                    logger.warning("Mammoth not available, skipping HTML conversion")

        except Exception as e:
            logger.warning(f"DOCX processing error: {e}")

        return tables

    def _process_image(self) -> List[Dict[str, Any]]:
        """Identify tables in image files."""
        logger.info(f"Identifying tables in image: {self.file_path}")

        tables = []
        try:
            img = cv2.imread(self.file_path)
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

            # Apply adaptive threshold
            thresh = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY_INV, 11, 3)

            # Line detection kernels
            horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (50, 1))
            vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 50))

            # Detect lines
            horizontal_lines = cv2.erode(thresh, horizontal_kernel, iterations=1)
            horizontal_lines = cv2.dilate(horizontal_lines, horizontal_kernel, iterations=3)

            vertical_lines = cv2.erode(thresh, vertical_kernel, iterations=1)
            vertical_lines = cv2.dilate(vertical_lines, vertical_kernel, iterations=3)

            # Combine lines
            table_mask = cv2.bitwise_or(horizontal_lines, vertical_lines)

            # Find contours
            contours, _ = cv2.findContours(table_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

            # Filter potential tables
            for contour in contours:
                area = cv2.contourArea(contour)
                if area > 5000:
                    x, y, w, h = cv2.boundingRect(contour)
                    if 0.5 < w / h < 5:
                        tables.append({
                            'method': 'image',
                            'bbox': (x, y, x + w, y + h),
                            'raw_image': img[y:y + h, x:x + w].copy()  # Ensure a copy
                        })
        except Exception as e:
            logger.warning(f"Image processing error: {e}")

        return tables


class ExtractionAgent:
    """Agent responsible for extracting data from identified tables."""

    def __init__(self, tables_info: List[Dict[str, Any]], file_path: str):
        self.tables_info = tables_info
        self.file_path = file_path
        self.extracted_tables = []

    def process(self) -> List[Dict[str, Any]]:
        """Extract data from each identified table."""
        logger.info(f"Extracting data from {len(self.tables_info)} tables")

        for i, table_info in enumerate(self.tables_info):
            method = table_info.get('method', '')

            try:
                table_data = []

                if method == 'tabula' or method == 'camelot':
                    df = table_info.get('data')
                    table_data = self._clean_dataframe(df)

                elif method == 'docx':
                    table_obj = table_info.get('table_obj')
                    table_data = self._extract_from_docx_table(table_obj)

                elif method == 'mammoth':
                    html = table_info.get('html')
                    table_data = self._extract_from_html(html)

                elif method == 'text':
                    # Process text lines into structured data
                    text_lines = table_info.get('text_content', [])
                    table_data = self._extract_from_text_lines(text_lines)

                elif method in ['visual', 'image', 'fullpage']:
                    img = table_info.get('raw_image')
                    if img is not None:
                        table_data = self._extract_with_ocr(img)

                else:
                    logger.warning(f"Unknown extraction method: {method}")
                    continue

                # Add to results if data was extracted
                if table_data and len(table_data) > 0:
                    extracted_table = {
                        'table_id': i + 1,
                        'method': method,
                        'data': table_data
                    }

                    # Add page number if available
                    if 'page' in table_info:
                        extracted_table['page'] = table_info['page']

                    self.extracted_tables.append(extracted_table)

            except Exception as e:
                logger.error(f"Error extracting table {i + 1}: {str(e)}")

        return self.extracted_tables

    def _clean_dataframe(self, df: pd.DataFrame) -> List[Dict[str, Any]]:
        """Clean and convert DataFrame to structured data."""
        try:
            # Handle DataFrame cleaning safely
            df = df.dropna(how='all').dropna(axis=1, how='all')
            df = df.fillna('')

            # Check for missing headers
            if all(isinstance(col, (int, float)) or str(col).startswith('Unnamed:') for col in df.columns):
                if not df.empty:
                    # Use first row as headers if available
                    first_row = df.iloc[0].tolist()
                    if any(str(x).strip() != '' for x in first_row):
                        headers = [str(x).strip() if str(x).strip() else f"Column_{i}" for i, x in enumerate(first_row)]
                        df.columns = headers
                        df = df.iloc[1:].reset_index(drop=True)
                    else:
                        df.columns = [f"Column_{i + 1}" for i in range(len(df.columns))]
                else:
                    df.columns = [f"Column_{i + 1}" for i in range(len(df.columns))]

            # Convert to list of dicts
            records = []
            for _, row in df.iterrows():
                # Convert each cell to string to avoid serialization issues
                record = {str(col): str(val) if val is not None else '' for col, val in row.items()}
                records.append(record)

            return records
        except Exception as e:
            logger.warning(f"Error cleaning DataFrame: {e}")
            return []

    def _extract_from_docx_table(self, table: DocxTable) -> List[Dict[str, Any]]:
        """Extract data from a DOCX table object."""
        try:
            data = []

            # Get headers from first row
            if len(table.rows) == 0:
                return []

            headers = []
            for cell in table.rows[0].cells:
                headers.append(cell.text.strip() or f"Column_{len(headers) + 1}")

            # Extract data rows
            for row in table.rows[1:]:
                row_data = {}
                for i, cell in enumerate(row.cells):
                    if i < len(headers):
                        row_data[headers[i]] = cell.text.strip()
                if row_data:
                    data.append(row_data)

            return data
        except Exception as e:
            logger.warning(f"Error extracting from DOCX table: {e}")
            return []

    def _extract_from_html(self, html: str) -> List[Dict[str, Any]]:
        """Extract table data from HTML content."""
        try:
            # Parse HTML tables
            tables = pd.read_html(html)

            all_data = []
            for df in tables:
                all_data.extend(self._clean_dataframe(df))

            return all_data
        except Exception as e:
            logger.warning(f"Error extracting from HTML: {e}")
            return []

    def _extract_from_text_lines(self, lines: List[str]) -> List[Dict[str, Any]]:
        """Extract table-like data from text lines."""
        if not lines or len(lines) < 2:
            return []

        try:
            # Look for delimiter patterns
            delimiters = ['|', '\t', '  ']
            best_delimiter = None
            max_columns = 0

            # Find the best delimiter
            for delimiter in delimiters:
                column_counts = []
                for line in lines:
                    parts = [part.strip() for part in line.split(delimiter) if part.strip()]
                    if len(parts) > 1:
                        column_counts.append(len(parts))

                if column_counts:
                    most_common_count = max(set(column_counts), key=column_counts.count)
                    if most_common_count > max_columns:
                        max_columns = most_common_count
                        best_delimiter = delimiter

            # If a good delimiter was found
            if best_delimiter and max_columns > 1:
                # Parse data rows
                data_rows = []
                for line in lines:
                    parts = [part.strip() for part in line.split(best_delimiter) if part.strip()]
                    if len(parts) > 1:
                        data_rows.append(parts)

                if not data_rows:
                    return []

                # Create structured data
                headers = data_rows[0] if len(data_rows) > 1 else [f"Column_{i + 1}" for i in range(len(data_rows[0]))]
                table_data = []

                for row in data_rows[1:] if len(data_rows) > 1 else data_rows:
                    row_dict = {}
                    for i, cell in enumerate(row):
                        if i < len(headers):
                            row_dict[headers[i]] = cell
                    if row_dict:
                        table_data.append(row_dict)

                return table_data

            # Fallback: Return lines as simple data
            return [{"Line_Content": line} for line in lines]

        except Exception as e:
            logger.warning(f"Error extracting from text lines: {e}")
            return []

    def _extract_with_ocr(self, img) -> List[Dict[str, Any]]:
        """Extract table data using OCR with improved error handling."""
        try:
            # Preprocess image
            if not isinstance(img, (np.ndarray, Image.Image)):
                logger.warning("Invalid image type for OCR")
                return []

            if isinstance(img, np.ndarray):
                # Convert BGR to RGB if needed
                if img.shape[2] == 3:
                    img_pil = Image.fromarray(cv2.cvtColor(img, cv2.COLOR_BGR2RGB))
                else:
                    img_pil = Image.fromarray(img)

                # Enhance contrast
                img_np = np.array(img_pil)
                img_np = cv2.convertScaleAbs(img_np, alpha=1.5, beta=0)
                img_pil = Image.fromarray(img_np)
            else:
                img_pil = img

            # Run OCR with error handling
            try:
                custom_config = r'--oem 3 --psm 6 -c preserve_interword_spaces=1'
                ocr_text = pytesseract.image_to_string(img_pil, config=custom_config)
            except Exception as e:
                logger.warning(f"OCR error: {e}")
                return []

            # Process the extracted text
            lines = [line.strip() for line in ocr_text.splitlines() if line.strip()]

            # Convert to structured data
            return self._extract_from_text_lines(lines)

        except Exception as e:
            logger.warning(f"Error in OCR extraction: {e}")
            return []


class ValidationAgent:
    """Agent responsible for validating and correcting extracted table data."""

    def __init__(self, tables: List[Dict[str, Any]], groq_api_key: Optional[str] = None):
        self.tables = tables
        self.groq_api_key = groq_api_key

    def process(self) -> List[Dict[str, Any]]:
        """Validate and correct the extracted tables."""
        logger.info(f"Validating and correcting {len(self.tables)} tables")
        validated_tables = []

        for table in self.tables:
            table_id = table.get('table_id', 0)
            data = table.get('data', [])

            if not data:
                logger.warning(f"Table {table_id} has no data to validate")
                continue

            # Basic validation and correction
            validation_issues = self._validate_table(data)

            if validation_issues:
                logger.info(f"Table {table_id} has {len(validation_issues)} issues, attempting correction")
                corrected_data = self._correct_table(data, validation_issues)
                table['data'] = corrected_data
                table['validation'] = 'corrected'
            else:
                table['validation'] = 'valid'

            validated_tables.append(table)

        return validated_tables

    def _validate_table(self, data: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Perform basic validation on table data."""
        issues = []

        if not data:
            return [{'type': 'empty_table', 'description': 'Table has no data'}]

        try:
            # Check for consistent columns
            all_keys = set()
            for row in data:
                all_keys.update(row.keys())

            for i, row in enumerate(data):
                row_keys = set(row.keys())
                missing_keys = all_keys - row_keys

                if missing_keys:
                    issues.append({
                        'type': 'missing_fields',
                        'row': i,
                        'fields': list(missing_keys)
                    })

            # Check for empty cells
            for i, row in enumerate(data):
                for key, value in row.items():
                    if value == '' or value is None:
                        issues.append({
                            'type': 'empty_cell',
                            'row': i,
                            'field': key
                        })
        except Exception as e:
            logger.warning(f"Error validating table: {e}")

        return issues

    def _correct_table(self, data: List[Dict[str, Any]], issues: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Attempt to correct table data issues."""
        try:
            corrected_data = data.copy()

            # Fix missing fields
            for issue in issues:
                if issue.get('type') == 'missing_fields':
                    row_idx = issue.get('row', 0)
                    fields = issue.get('fields', [])

                    if 0 <= row_idx < len(corrected_data):
                        for field in fields:
                            corrected_data[row_idx][field] = ''

                elif issue.get('type') == 'empty_cell':
                    row_idx = issue.get('row', 0)
                    field = issue.get('field', '')

                    if 0 <= row_idx < len(corrected_data) and field in corrected_data[row_idx]:
                        # Try to infer a value from other rows
                        values = [row.get(field, '') for i, row in enumerate(data)
                                  if i != row_idx and field in row and row[field]]

                        if values:
                            # Use most common value
                            from collections import Counter
                            most_common = Counter(values).most_common(1)
                            if most_common:
                                corrected_data[row_idx][field] = most_common[0][0]

            # Use LLM if available and table has significant data
            if self.groq_api_key and client and len(data) > 3:
                try:
                    # Sample data for LLM correction
                    sample = data[:3]
                    corrected_sample = self._llm_correction(sample)

                    # Apply corrections to sample
                    if corrected_sample and len(corrected_sample) == len(sample):
                        for i in range(len(sample)):
                            corrected_data[i] = corrected_sample[i]
                except Exception as e:
                    logger.warning(f"LLM correction error: {e}")

            return corrected_data

        except Exception as e:
            logger.warning(f"Error correcting table: {e}")
            return data

    def _llm_correction(self, sample_data: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Use LLM to correct a sample of table data."""
        if not self.groq_api_key or not client or not sample_data:
            return None

        try:
            # Create a safe JSON string
            safe_json = json.dumps(sample_data, ensure_ascii=True)

            prompt = f"""
            Review and correct this table data sample:
            ```
            {safe_json}
            ```
            Fix any inconsistencies or errors. Return only the corrected JSON with no extra text.
            """

            # Get LLM response
            response = client.chat.completions.create(
                messages=[
                    {"role": "system",
                     "content": "You are a data validation assistant that fixes errors in table data. Return only valid JSON."},
                    {"role": "user", "content": prompt}
                ],
                model="llama-3.3-70b-versatile",
                max_tokens=1000,
                temperature=0.1
            )

            # Extract and parse the JSON response
            result = response.choices[0].message.content.strip()
            time.sleep(10)
            # Remove code block syntax if present
            if result.startswith("```") and result.endswith("```"):
                result = result[3:-3].strip()
            if result.startswith("```json") and result.endswith("```"):
                result = result[7:-3].strip()

            try:
                corrected_data = json.loads(result)
                if isinstance(corrected_data, list):
                    return corrected_data
            except json.JSONDecodeError:
                logger.warning("Failed to parse LLM response as JSON")

        except Exception as e:
            logger.warning(f"LLM processing error: {e}")

        return None


class TableExtractionPipeline:
    """Main pipeline coordinating the table extraction process."""

    def __init__(self, file_path: str, output_path: str = None, groq_api_key: str = None):
        self.file_path = file_path
        self.output_path = output_path or f"{os.path.splitext(file_path)[0]}_tables.json"
        self.groq_api_key = groq_api_key

    def run(self) -> Dict[str, Any]:
        """Execute the full table extraction pipeline."""
        logger.info(f"Starting table extraction for {self.file_path}")

        try:
            # Step 1: Identify tables
            document_agent = DocumentAgent(self.file_path)
            tables_info = document_agent.process()
            logger.info(f"Found {len(tables_info)} potential tables")

            if not tables_info:
                logger.warning("No tables found in the document")
                result = {"tables": [], "message": "No tables found"}
            else:
                # Step 2: Extract data
                extraction_agent = ExtractionAgent(tables_info, self.file_path)
                extracted_tables = extraction_agent.process()
                logger.info(f"Extracted data from {len(extracted_tables)} tables")

                # Step 3: Validate and correct
                validation_agent = ValidationAgent(extracted_tables, self.groq_api_key)
                validated_tables = validation_agent.process()

                result = {
                    "file": os.path.basename(self.file_path),
                    "tables": validated_tables,
                    "table_count": len(validated_tables),
                    "extraction_timestamp": str(pd.Timestamp.now())
                }

            # Save results
            try:
                with open(self.output_path, 'w', encoding='utf-8') as f:
                    json.dump(result, f, indent=2, ensure_ascii=False)
            except:
                print("Not possible to get the tables inside")
        except:
            print("Not possible to get the tables inside")
        logger.info(f"Saved extraction results to {self.output_path}")

        return result


def main():


    file_path = r"C:\Users\SanskarZanwar\PycharmProjects\TableDataExctraction\input\SYB67_1_202411_Population, Surface Area and Density.pdf"
    output_path = r"C:\Users\SanskarZanwar\PycharmProjects\TableDataExctraction\output"
    if not os.path.isfile(file_path):
        print(f"Error: Input file path not found.")
        return
    output_path = os.path.join(output_path, "extracted_tables.json")
    # Run the pipeline
    pipeline = TableExtractionPipeline(
        file_path=file_path,
        output_path=output_path,
        groq_api_key=GROQ_API_KEY
    )
    result = pipeline.run()

    print(f"\nExtraction complete!")
    #print(f"Found {result['table_count']} tables in {os.path.basename(file_path)}")
    print(f"Results saved to {pipeline.output_path}")


if __name__ == "__main__":
    main()